from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "work_vkr"
OUT.mkdir(exist_ok=True)


def extract_docx(path: Path, out: Path) -> None:
    doc = Document(path)
    lines: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table_index, table in enumerate(doc.tables, start=1):
        lines.append(f"\n[Таблица {table_index}]")
        for row in table.rows:
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    out.write_text("\n".join(lines), encoding="utf-8")


def extract_pdf(path: Path, out: Path) -> None:
    reader = PdfReader(str(path))
    chunks: list[str] = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        chunks.append(f"\n--- PAGE {page_num} ---\n{text}")
    out.write_text("\n".join(chunks), encoding="utf-8")


if __name__ == "__main__":
    extract_docx(Path(r"D:\Desktop\ВКР.docx"), OUT / "vkr_extracted.txt")
    extract_pdf(Path(r"C:\Users\Den4ik\Downloads\sfu-stu-7.5-07 (1).pdf"), OUT / "stu_extracted.txt")
